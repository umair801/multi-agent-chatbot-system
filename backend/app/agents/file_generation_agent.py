import logging
import os
import io
from typing import Optional, Dict, Any, List
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class FileGenerationAgent:
    """
    Specialist agent for generating PDF, DOCX, and CSV files
    """
    
    def __init__(self):
        self.output_dir = '/tmp/agent_outputs'
        os.makedirs(self.output_dir, exist_ok=True)
    
    async def generate_csv(self, data: List[Dict[str, Any]], filename: str) -> dict:
        """
        Generate a CSV file from data
        
        Args:
            data: List of dictionaries (each dict is a row)
            filename: Name of the output file (without extension)
        
        Returns:
        {
            'status': 'success' | 'failed',
            'file_path': str,
            'file_name': str,
            'file_size': int,
            'error': str (optional)
        }
        """
        try:
            logger.info(f'Generating CSV: {filename}')
            
            # Convert to DataFrame
            df = pd.DataFrame(data)
            
            # Generate filename with timestamp
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            file_name = f'{filename}_{timestamp}.csv'
            file_path = os.path.join(self.output_dir, file_name)
            
            # Write to CSV
            df.to_csv(file_path, index=False)
            
            file_size = os.path.getsize(file_path)
            logger.info(f'CSV generated: {file_name} ({file_size} bytes)')
            
            return {
                'status': 'success',
                'file_path': file_path,
                'file_name': file_name,
                'file_size': file_size,
                'output_type': 'csv'
            }
        
        except Exception as e:
            logger.error(f'CSV generation failed: {str(e)}')
            return {
                'status': 'failed',
                'error': str(e)
            }
    
    async def generate_docx(self, title: str, content: str, filename: str) -> dict:
        """
        Generate a DOCX file
        
        Args:
            title: Document title
            content: Document body text (can include newlines)
            filename: Name of the output file (without extension)
        
        Returns:
        {
            'status': 'success' | 'failed',
            'file_path': str,
            'file_name': str,
            'file_size': int,
            'error': str (optional)
        }
        """
        try:
            logger.info(f'Generating DOCX: {filename}')
            
            # Create document
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content
            for paragraph_text in content.split('\n\n'):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            
            # Generate filename with timestamp
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            file_name = f'{filename}_{timestamp}.docx'
            file_path = os.path.join(self.output_dir, file_name)
            
            # Save document
            doc.save(file_path)
            
            file_size = os.path.getsize(file_path)
            logger.info(f'DOCX generated: {file_name} ({file_size} bytes)')
            
            return {
                'status': 'success',
                'file_path': file_path,
                'file_name': file_name,
                'file_size': file_size,
                'output_type': 'docx'
            }
        
        except Exception as e:
            logger.error(f'DOCX generation failed: {str(e)}')
            return {
                'status': 'failed',
                'error': str(e)
            }
    
    async def generate_pdf(self, title: str, content: str, filename: str) -> dict:
        """
        Generate a PDF file using WeasyPrint
        
        Args:
            title: Document title
            content: Document body text (can include newlines)
            filename: Name of the output file (without extension)
        
        Returns:
        {
            'status': 'success' | 'failed',
            'file_path': str,
            'file_name': str,
            'file_size': int,
            'error': str (optional)
        }
        """
        try:
            logger.info(f'Generating PDF: {filename}')
            
            from weasyprint import HTML, CSS
            
            # Generate filename with timestamp
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            file_name = f'{filename}_{timestamp}.pdf'
            file_path = os.path.join(self.output_dir, file_name)
            
            # Create HTML content
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 2cm; }}
                    h1 {{ text-align: center; color: #333; }}
                    p {{ line-height: 1.6; color: #555; }}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                <div>{''.join(f'<p>{para}</p>' for para in content.split(chr(10)+chr(10)) if para.strip())}</div>
            </body>
            </html>
            """
            
            # Generate PDF
            HTML(string=html_content).write_pdf(file_path)
            
            file_size = os.path.getsize(file_path)
            logger.info(f'PDF generated: {file_name} ({file_size} bytes)')
            
            return {
                'status': 'success',
                'file_path': file_path,
                'file_name': file_name,
                'file_size': file_size,
                'output_type': 'pdf'
            }
        
        except Exception as e:
            logger.error(f'PDF generation failed: {str(e)}')
            return {
                'status': 'failed',
                'error': str(e)
            }
    
    async def generate(self, output_type: str, **kwargs) -> dict:
        """
        Generic file generation dispatcher
        
        Args:
            output_type: 'pdf', 'docx', or 'csv'
            **kwargs: Type-specific arguments
        
        Returns: File generation result dict
        """
        if output_type == 'csv':
            return await self.generate_csv(
                data=kwargs.get('data', []),
                filename=kwargs.get('filename', 'output')
            )
        elif output_type == 'docx':
            return await self.generate_docx(
                title=kwargs.get('title', 'Document'),
                content=kwargs.get('content', ''),
                filename=kwargs.get('filename', 'output')
            )
        elif output_type == 'pdf':
            return await self.generate_pdf(
                title=kwargs.get('title', 'Document'),
                content=kwargs.get('content', ''),
                filename=kwargs.get('filename', 'output')
            )
        else:
            logger.error(f'Unknown output type: {output_type}')
            return {
                'status': 'failed',
                'error': f'Unsupported output type: {output_type}'
            }

# Global file generation agent instance
file_generation_agent = FileGenerationAgent()
